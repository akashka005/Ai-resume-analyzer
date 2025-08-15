from docx import Document
from docx.shared import Pt

def generate_final_docx(contact, education, experience, skills, rewrites, filename='optimized_resume.docx'):
    doc = Document()
    name = contact.get('name') or "Candidate Name"
    doc.add_heading(name, level=0)
    contact_line = ", ".join(filter(None,[contact.get('email'), contact.get('phone'), contact.get('linkedin')]))
    doc.add_paragraph(contact_line)
    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph("Optimized professional summary: " + (rewrites.get('summary_suggestion') or ""))
    doc.add_heading('Skills', level=1)
    skills_par = doc.add_paragraph()
    skills_par.add_run(", ".join(skills.get('technical', []) + skills.get('soft', []))) 
    doc.add_heading('Experience', level=1)
    for p in experience.get('positions', []):
        ptitle = p.get('title') or "Title"
        company = p.get('company') or ""
        para = doc.add_paragraph()
        para.add_run(f"{ptitle} — {company}").bold = True
        para.add_run("\n" + p.get('raw',''))
    doc.add_heading('Education', level=1)
    for e in education.get('degrees', []):
        doc.add_paragraph(e.get('line',''))
    doc.add_heading('Suggestions & Rewrites', level=1)
    for r in rewrites.get('experience_rewrites', []):
        doc.add_paragraph("Original: " + r['original'])
        doc.add_paragraph("Suggested: " + r['suggested'])

    doc.save(filename)
    return filename